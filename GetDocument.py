from docx import Document
from docx.shared import Pt
from expression import Expressions


def get_document(exps : Expressions, file_name : str , with_ans : bool = False) -> None:
    document = Document()
        
    style = document.styles['Normal']
    font = style.font
    # font.name = 'Arial'
    font.size = Pt(20)

    document.add_heading('Test', 0)
    document.add_heading('                                                                              Namn: ____________________', 1)

    table = document.add_table(rows=0, cols=exps.ncol)

    for i in range(exps.nrow):

        row_cells = table.add_row().cells

        for j in range(exps.ncol):
            if with_ans:
                cell = exps.expressions[j + exps.ncol * i].answer
            else:
                cell = exps.expressions[j + exps.ncol * i].txt
                
            row_cells[j].text = cell

    document.save(f'{file_name}')