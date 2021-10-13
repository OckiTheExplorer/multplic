from docx import Document
from docx.shared import Pt
from expression import Expressions
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_test(exps : Expressions, with_ans : bool = False) -> None:
    document = Document()
        
    style = document.styles['Normal']
    font = style.font
    # font.name = 'Arial'
    font.size = Pt(20)

    document.add_heading('Test', 0)
    header = document.add_heading('Namn: ____________________', 1)
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    table = document.add_table(rows=0, cols=exps.ncol)

    for i in range(exps.nrow):

        row_cells = table.add_row().cells

        for j in range(exps.ncol):
            if with_ans:
                cell = exps.expressions[j + exps.ncol * i].answer
            else:
                cell = exps.expressions[j + exps.ncol * i].txt
                
            row_cells[j].text = cell

    return document
